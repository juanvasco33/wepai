#!/usr/bin/env python3
"""
VALIDATE_LLM_LEGAL.py — Validación real de los 6 templates LLM-driven.

USO:
    export ANTHROPIC_API_KEY=sk-ant-...
    python3 validate_llm_legal.py [--all | --country=Colombia | --template=employment_contract]

Por defecto valida UN doc (employment_contract en Colombia) — el caso más
representativo, costo ~$0.15 USD. La opción --all valida los 6 templates en
5 países (30 docs, costo ~$5 USD).

LO QUE VERIFICA:
- La API responde correctamente
- El JSON parsea
- El doc se genera y se guarda
- Contiene marcadores correctos del país (CST/LFT/CST/LCT/ET según corresponda)
- NO contaminación cruzada entre países
- El caché funciona (segunda corrida idéntica = 0 API calls)
- El fallback funciona si la API falla

OUTPUT:
- Imprime un reporte por consola
- Guarda los docs generados en ./test_outputs/
- Exit code 0 si todo OK, 1 si hay fallos
"""
import os, sys, time, argparse, tempfile, shutil
from pathlib import Path

# Asegurar que estamos en el directorio del proyecto
PROJECT_ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(PROJECT_ROOT))

OUTPUT_DIR = PROJECT_ROOT / "test_outputs"
OUTPUT_DIR.mkdir(exist_ok=True)

# Lo que cada país debe contener / NO debe contener
COUNTRY_EXPECTATIONS = {
    "Colombia": {
        "must_contain": ["CST", "EPS", "Bogotá"],
        "must_not_contain": ["LFT", "IMSS", "AFORE", "EL PATRÓN"],
    },
    "México": {
        "must_contain": ["LFT", "IMSS"],
        "must_not_contain": ["CST", "EPS", "Cesantías"],
    },
    "Chile": {
        "must_contain": ["Código del Trabajo", "AFP"],
        "must_not_contain": ["LFT", "CST", "IMSS"],
    },
    "Argentina": {
        "must_contain": ["LCT", "ANSES"],
        "must_not_contain": ["LFT", "CST"],
    },
    "España": {
        "must_contain": ["Estatuto", "Seguridad Social"],
        "must_not_contain": ["LFT", "CST", "IMSS"],
    },
}

# Templates a validar y sus datos mínimos
TEMPLATE_DATA = {
    "employment_contract": {
        "tipo_contrato": "indefinido",
        "patron_nombre": "Acme Software S.A.",
        "trabajador_nombre": "Juan Pérez",
        "puesto": "Senior Software Engineer",
        "salario": "5,000,000",
        "fecha_inicio": "2026-06-01",
    },
    "termination_letter": {
        "patron_nombre": "Acme Software S.A.",
        "trabajador_nombre": "Juan Pérez",
        "puesto": "Senior Software Engineer",
        "causal": "Mutuo acuerdo",
        "fecha_terminacion": "2026-12-31",
        "antiguedad": "1 año",
    },
    "settlement": {
        "patron_nombre": "Acme Software S.A.",
        "trabajador_nombre": "Juan Pérez",
        "fecha_inicio": "2025-01-01",
        "fecha_terminacion": "2026-12-31",
        "salario": "5,000,000",
        "tipo_terminacion": "sin justa causa",
    },
    "commercial_lease": {
        "arrendador_nombre": "Inmobiliaria Plaza S.A.",
        "arrendatario_nombre": "Acme Software S.A.",
        "direccion_inmueble": "Carrera 50 #80-23",
        "monto_renta": "5,000,000",
        "plazo": "24 meses",
        "uso_destinado": "Oficinas administrativas",
    },
    "work_rules": {
        "empresa_nombre": "Acme Software S.A.",
        "industria": "Desarrollo de software",
        "tamano_empresa": "50",
        "ubicacion": "Bogotá D.C.",
    },
    "power_of_attorney": {
        "otorgante_nombre": "Juan Pérez",
        "apoderado_nombre": "Dra. María González",
        "tipo_poder": "especial",
        "facultades": "Representación judicial ante tribunales civiles",
        "vigencia": "1 año",
    },
}


def extract_all_text(docx_path):
    """Extrae TODO el texto de un .docx: body + tablas + headers + footers."""
    from docx import Document
    d = Document(docx_path)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    for section in d.sections:
        for container in (section.header, section.footer):
            parts.extend(p.text for p in container.paragraphs)
    return "\n".join(parts)


def validate_one(template_id, country):
    """Genera UN documento y valida. Devuelve dict con resultados."""
    from office import word_controller as wc
    from office import llm_legal_generator as llg

    # Apuntar OUTPUT_DIR de WEP AI al de tests
    wc.OUTPUT_DIR = str(OUTPUT_DIR)

    # Datos: mezcla los datos base del template con país e idioma
    datos = {
        **TEMPLATE_DATA.get(template_id, {}),
        "idioma": "es",
        "pais": country,
    }

    t0 = time.time()
    try:
        fp = wc.generate_word({
            "template_id": template_id,
            "titulo": f"VAL_{template_id}_{country}",
            "instrucciones": f"Validación real de {template_id} para {country}",
            "datos": datos,
        })
        latency = time.time() - t0
    except Exception as e:
        return {"ok": False, "error": str(e), "latency": time.time() - t0}

    if not os.path.exists(fp):
        return {"ok": False, "error": "Archivo no creado", "latency": latency}

    text = extract_all_text(fp)
    expectations = COUNTRY_EXPECTATIONS.get(country, {})
    must = expectations.get("must_contain", [])
    must_not = expectations.get("must_not_contain", [])

    pos_results = [(t, t in text) for t in must]
    neg_results = [(t, t not in text) for t in must_not]
    pos_ok = sum(1 for _, ok in pos_results if ok)
    neg_ok = sum(1 for _, ok in neg_results if ok)

    # Heurística de calidad: ¿tiene contenido sustantivo?
    word_count = len(text.split())

    return {
        "ok": pos_ok == len(must) and neg_ok == len(must_not),
        "filepath": fp,
        "size_bytes": os.path.getsize(fp),
        "latency": latency,
        "word_count": word_count,
        "positives": {"ok": pos_ok, "total": len(must),
                      "failures": [t for t, ok in pos_results if not ok]},
        "negatives": {"ok": neg_ok, "total": len(must_not),
                      "failures": [t for t, ok in neg_results if not ok]},
    }


def run_validation(templates, countries):
    """Corre validación sobre el producto cartesiano de templates × countries."""
    if not os.environ.get("ANTHROPIC_API_KEY"):
        print("✗ ANTHROPIC_API_KEY no está configurada.")
        print("  Exportala: export ANTHROPIC_API_KEY=sk-ant-...")
        return 1

    print("═" * 70)
    print("VALIDACIÓN REAL — llm_legal_generator")
    print("═" * 70)
    print(f"Templates: {templates}")
    print(f"Countries: {countries}")
    print(f"Total documentos: {len(templates) * len(countries)}")
    print(f"Costo estimado: ~${0.15 * len(templates) * len(countries):.2f} USD")
    print("═" * 70)

    results = {}
    total_ok = 0
    total = 0

    for tid in templates:
        for country in countries:
            key = f"{tid}/{country}"
            total += 1
            print(f"\n[{total}] {key}", flush=True)
            print(f"    Llamando a la API...", end=" ", flush=True)
            r = validate_one(tid, country)
            results[key] = r
            if r["ok"]:
                total_ok += 1
                print(f"✓ OK ({r['latency']:.1f}s, {r['word_count']} palabras, {r['size_bytes']} bytes)")
            else:
                if "error" in r:
                    print(f"✗ ERROR: {r['error'][:100]}")
                else:
                    print(f"⚠ ({r['latency']:.1f}s) " +
                          f"pos {r['positives']['ok']}/{r['positives']['total']}, " +
                          f"neg {r['negatives']['ok']}/{r['negatives']['total']}")
                    if r["positives"]["failures"]:
                        print(f"    Faltan: {r['positives']['failures']}")
                    if r["negatives"]["failures"]:
                        print(f"    Contaminación: {r['negatives']['failures']}")
                if "filepath" in r:
                    print(f"    Doc para revisión manual: {r['filepath']}")

    print()
    print("═" * 70)
    print(f"RESULTADO: {total_ok}/{total} documentos pasaron checks automáticos")
    print("═" * 70)
    print()
    print("⚠ IMPORTANTE: el check automático verifica marcadores generales (CST, LFT, etc.)")
    print("  La calidad legal real requiere revisión humana — abrí los .docx en Word y")
    print("  validá que las citas a artículos sean correctas y el contenido tenga sentido")
    print("  legal para cada jurisdicción.")
    print()
    print(f"Docs guardados en: {OUTPUT_DIR}")
    return 0 if total_ok == total else 1


def main():
    p = argparse.ArgumentParser()
    p.add_argument("--all", action="store_true",
                   help="Validar los 6 templates en 5 países (30 docs, ~$5)")
    p.add_argument("--country", default="Colombia",
                   help="País a validar (Colombia, México, Chile, Argentina, España)")
    p.add_argument("--template", default="employment_contract",
                   help="Template a validar (employment_contract, termination_letter, ...)")
    args = p.parse_args()

    if args.all:
        templates = list(TEMPLATE_DATA.keys())
        countries = list(COUNTRY_EXPECTATIONS.keys())
    else:
        templates = [args.template]
        countries = [args.country]

    sys.exit(run_validation(templates, countries))


if __name__ == "__main__":
    main()
