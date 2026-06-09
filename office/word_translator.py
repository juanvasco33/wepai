"""
WEP AI — Capa de traducción Word ES → EN (v15.0)

Estrategia: en lugar de reescribir 38 builders para que soporten EN de forma
nativa, dejamos que los builders generen en español (como hoy) y traducimos
el .docx resultante usando Claude. Un caché JSON crece con cada documento y
amortiza el costo: frases legales repetidas (cláusulas estándar, disclaimers,
labels de tablas) solo se traducen una vez en la vida del producto.

Integración en generate_word():

    if lang == "en":
        from office.word_translator import translate_doc_to_en
        doc = translate_doc_to_en(doc, doc_type=template_id, cfg=cfg)

Costo aproximado:
  - Primer documento de un tipo nuevo:    ~$0.05–0.10 USD (Sonnet)
  - Documentos posteriores del mismo tipo: ~$0.00–0.02 (mayoría cacheada)

Caché:
  - Ruta: ~/.wepai/word_translations_cache.json
  - Schema: dict[str, str]  (texto_es → texto_en)
  - Cargado lazy en la primera llamada; persistido tras cada batch exitoso.

Fallback (sin red, sin API key, JSON malformado, etc.):
  - Devuelve el documento en español original
  - Inserta un banner rojo arriba del doc avisando del fallo
  - No lanza excepción — el flujo de generate_word sigue normalmente
"""

import os
import re
import json
import logging
import threading
from typing import Iterable

import anthropic

log = logging.getLogger("wepai.word_translator")

# ── CONFIGURACIÓN ─────────────────────────────────────────────────────────────
CACHE_PATH = os.path.expanduser("~/.wepai/word_translations_cache.json")
# v15.4: model ID centralizado en config.py (override por WEPAI_MODEL_TRANSLATOR)
from config import DEFAULT_TRANSLATOR_MODEL as MODEL
MAX_BATCH_SIZE = 60  # párrafos por API call

# Cliente Anthropic con la misma config que agent/brain.py
_client = anthropic.Anthropic(
    api_key=os.environ.get("ANTHROPIC_API_KEY", ""),
    max_retries=3,
    timeout=60.0,
)

_cache: dict[str, str] = {}
_cache_lock = threading.Lock()
_cache_loaded = False


# ══════════════════════════════════════════════════════════════════════════════
# SYSTEM PROMPT
# ══════════════════════════════════════════════════════════════════════════════
# Diseñado específicamente para documentos legales/business ES → EN.
# Cada regla aquí está peleada con casos reales del codebase de WEP AI:
# - "rescisión" → "termination" (NO "rescission" — concepto distinto)
# - Mantener "Art. 50 LFT" verbatim porque es un identificador legal
# - Conservar placeholders [Nombre] pero traduciéndolos a [Name]
# - Preservar nombres propios, IDs fiscales, montos sin tocar
_SYSTEM_PROMPT = """You translate fragments of a legal/business document from Spanish to English.
The document type is "{doc_type}" and the country context is "{country}".

STRICT RULES:

1. Translate Spanish → English. If a fragment is ALREADY in English, return it unchanged (do not "correct" it).

2. NEVER translate — preserve exactly as input:
   - Numbers, percentages, currency: "$50,000 MXN", "16%", "5,000.00"
   - Proper names of people: "Juan Pérez", "María García López"
   - Company / legal entity names: "Tigo S.A.", "Logística XYZ S.A.S."
   - Tax/ID numbers: "RFC ABC123456", "NIT 900123456-7", "CC 12345678", "CUIT 30-12345678-9"
   - Email addresses, URLs, phone numbers
   - Symbols: ⚠ ✓ ✗ • — · «» " "
   - Legal article references stay verbatim, with optional clarification in parens added ONCE:
     "Art. 50 LFT" → "Art. 50 LFT (Mexican Federal Labor Law)"
     "Artículo 1764 del Código Civil" → "Article 1764 of the Civil Code"
     Never invent translations of the article TEXT — only annotate the citation.

3. PLACEHOLDERS in square brackets: translate the GENERIC label, keep the brackets.
   "[Nombre]" → "[Name]"
   "[Dirección del Prestador]" → "[Provider's Address]"
   "[Razón Social del Cliente]" → "[Client's Legal Name]"
   "[N] meses" → "[N] months"
   BUT if the bracket content is clearly user-specific data (full address, full name), keep it as-is.

4. LEGAL TERMINOLOGY — use the correct English legal equivalent, not a literal translation:
   - "rescisión" / "rescindir" → "termination" / "terminate" (NOT "rescission" — different concept in common law)
   - "objeto" (in a contract clause) → "subject matter" or "scope" (per context)
   - "el patrón" / "el empleador" → "the employer"
   - "el trabajador" → "the employee"
   - "el prestador" / "prestador de servicios" → "the service provider"
   - "el cliente" → "the client"
   - "el otorgante" → "the grantor"; "el apoderado" → "the attorney-in-fact"
   - "el cedente" / "el cesionario" → "the assignor" / "the assignee"
   - "el arrendador" / "el arrendatario" → "the landlord" / "the tenant" (real estate) or "the lessor" / "the lessee" (general)
   - "el vendedor" / "el comprador" → "the seller" / "the buyer"
   - "caso fortuito" / "fuerza mayor" → "force majeure"
   - "domicilio" → "registered address" (legal) or "address" (general)
   - "fuero" (renunciar al fuero) → "waive any jurisdictional defense"
   - "comparecen" → "the parties hereby enter into" / "the parties appear"
   - "se obliga a" → "agrees to" / "undertakes to"
   - "aguinaldo" → "Christmas bonus / 13th-month pay" (preserve specificity)
   - "prima de servicios" (Colombia) → "service premium"
   - "indemnización" → "severance" (employment) or "indemnification" (general)

5. ORDINAL CLAUSE NUMBERS (very common): convert one-to-one.
   PRIMERA→FIRST · SEGUNDA→SECOND · TERCERA→THIRD · CUARTA→FOURTH · QUINTA→FIFTH
   SEXTA→SIXTH · SÉPTIMA→SEVENTH · OCTAVA→EIGHTH · NOVENA→NINTH · DÉCIMA→TENTH
   DÉCIMA PRIMERA→ELEVENTH · DÉCIMA SEGUNDA→TWELFTH

6. DATES: "21 de mayo de 2026" → "May 21, 2026"
   Months: enero→January, febrero→February, marzo→March, abril→April, mayo→May,
           junio→June, julio→July, agosto→August, septiembre→September,
           octubre→October, noviembre→November, diciembre→December.

7. PRESERVE STRUCTURE:
   - Line breaks (\\n) stay as line breaks; do NOT collapse paragraphs into one line
   - Bullet/list markers stay: "a)", "b)", "•", "—", "1.", "i)"
   - Capitalization of section titles is preserved (UPPERCASE stays UPPERCASE, Title Case stays Title Case)
   - Trailing/leading whitespace is preserved

OUTPUT FORMAT — respond with ONLY a single valid JSON object, no prose, no markdown fences:

{{"items": [{{"id": "<id_from_input>", "text": "<translated_text>"}}, ...]}}

One output item per input item, in the same order, every input "id" matched."""


# ══════════════════════════════════════════════════════════════════════════════
# CACHÉ EN DISCO
# ══════════════════════════════════════════════════════════════════════════════
def _load_cache():
    """Carga el caché desde disco una sola vez (thread-safe, idempotente)."""
    global _cache, _cache_loaded
    with _cache_lock:
        if _cache_loaded:
            return
        _cache_loaded = True
        try:
            if os.path.exists(CACHE_PATH):
                with open(CACHE_PATH, "r", encoding="utf-8") as f:
                    _cache = json.load(f)
                log.info("Caché cargado: %d entradas desde %s", len(_cache), CACHE_PATH)
            else:
                _cache = {}
        except Exception as e:
            log.warning("No se pudo cargar el caché de traducciones: %s", e)
            _cache = {}


def _save_cache():
    """Persiste el caché a disco. Llamado tras cada batch exitoso."""
    with _cache_lock:
        try:
            os.makedirs(os.path.dirname(CACHE_PATH), exist_ok=True)
            with open(CACHE_PATH, "w", encoding="utf-8") as f:
                json.dump(_cache, f, ensure_ascii=False, indent=2, sort_keys=True)
        except Exception as e:
            log.warning("No se pudo guardar el caché: %s", e)


# ══════════════════════════════════════════════════════════════════════════════
# DETECCIÓN DE TEXTO TRADUCIBLE
# ══════════════════════════════════════════════════════════════════════════════
# Texto puramente numérico/simbólico no necesita ir al traductor.
# Acepta también códigos de moneda comunes después del número (MXN, USD, etc.)
_SKIP_PATTERN = re.compile(
    r'^[\d\s\$£€¥,.\-:;%/\\()|+*°#&@]+(\s*(MXN|USD|EUR|COP|ARS|CLP|PEN|BRL|GBP|JPY|CAD|CHF|MN))?$'
)

# Patrones que indican "claramente inglés" — saltarlos ahorra API calls.
# Heurística conservadora: si el texto tiene palabras 100% inglesas claras
# Y NO tiene ningún indicador español, saltarlo.
_ENGLISH_ONLY_HINTS = re.compile(
    r'\b(the|and|of|to|for|with|that|this|shall|hereby|whereas|'
    r'agreement|party|parties|services?|employee|employer|client|'
    r'provider|company|herein|thereof|signature|signed|page)\b',
    re.IGNORECASE,
)

# Cualquier indicio de español — tildes, ñ, signos invertidos, o palabras
# claramente españolas. Si alguno aparece → traducir.
_SPANISH_LETTERS = set("ñáéíóúüÑÁÉÍÓÚÜ¿¡")

_SPANISH_HINTS = re.compile(
    r'\b(que|de|la|el|los|las|del|al|para|con|por|sin|una|uno|sus?|'
    r'es|son|sera|este|esta|presente|partes?|contrato|acuerdo|clausula|'
    r'empresa|empleado|trabajador|patron|cliente|prestador|arrendador|'
    r'comparecen|obliga|compromete|segun|conforme|primera|segunda|tercera|'
    r'cuarta|quinta|sexta|septima|octava|novena|decima|objeto|monto|'
    r'duracion|salario|campo|valor|nombre|direccion|ciudad|fecha|'
    r'firma|reciben|otorgan|certifican)\b',
    re.IGNORECASE,
)


def _should_translate(text: str) -> bool:
    """True si el fragmento debería ir al traductor.

    Estrategia: permisiva por defecto. Solo saltamos lo claramente innecesario
    (vacío, demasiado corto, solo símbolos/números, o claramente inglés sin
    ningún rastro de español). Cuando hay duda, mandamos al LLM — el system
    prompt le instruye a devolver texto inglés sin cambios, y el caché
    amortiza el costo de los repetidos.
    """
    if not text:
        return False
    t = text.strip()
    if len(t) < 2:
        return False
    if _SKIP_PATTERN.match(t):
        return False
    # Si tiene menos de 3 caracteres alfabéticos (descontando símbolos y
    # números), no hay nada significativo que traducir — saltar.
    alpha_only = re.sub(r'[^a-zA-ZñáéíóúüÑÁÉÍÓÚÜ]', '', t)
    if len(alpha_only) < 3:
        return False
    # Si tiene marcadores ortográficos españoles → traducir sí o sí
    if any(c in t for c in _SPANISH_LETTERS):
        return True
    # Si tiene palabras claramente españolas → traducir
    if _SPANISH_HINTS.search(t):
        return True
    # Si tiene palabras claramente inglesas Y nada español → saltar
    if _ENGLISH_ONLY_HINTS.search(t):
        return False
    # Caso ambiguo (nombre propio, label corto sin marcadores): por las dudas, traducir.
    # El LLM lo pasará tal cual si es nombre, y el caché lo absorbe.
    return True


def _iter_translatable_units(doc) -> Iterable[tuple[str, object]]:
    """Itera todas las unidades traducibles del documento.

    Cubre: párrafos del body, celdas de tablas, headers y footers de cada sección.
    Yield (location_id, paragraph_obj). El location_id es solo para mapeo
    bidireccional con la respuesta del LLM y para logs.
    """
    # Body
    for i, p in enumerate(doc.paragraphs):
        if p.text and _should_translate(p.text):
            yield (f"b{i}", p)
    # Tablas
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    if p.text and _should_translate(p.text):
                        yield (f"t{ti}r{ri}c{ci}p{pi}", p)
    # Headers y footers (por sección)
    for si, section in enumerate(doc.sections):
        for kind, container in (("hdr", section.header), ("ftr", section.footer)):
            for pi, p in enumerate(container.paragraphs):
                if p.text and _should_translate(p.text):
                    yield (f"s{si}{kind}{pi}", p)


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    """Sustituye el texto del párrafo preservando el formato del primer run.

    Limitación conocida: párrafos con formato mixto (negrita-luego-regular)
    pierden los bordes internos — todo el texto traducido hereda el formato
    del primer run. Para el output típico de los builders de WEP AI (un solo
    run por párrafo, generado vía `_para()`), esto es correcto.
    """
    if not paragraph.runs:
        return
    first = paragraph.runs[0]
    first.text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _chunks(seq, n):
    """Divide una lista en chunks de tamaño n."""
    for i in range(0, len(seq), n):
        yield seq[i:i + n]


# ══════════════════════════════════════════════════════════════════════════════
# LLAMADA AL LLM
# ══════════════════════════════════════════════════════════════════════════════
def _llm_translate(items: list[tuple[str, str]], doc_type: str, country: str) -> dict[str, str]:
    """Manda un batch a Claude. Devuelve {id: traducción} o {} si falla."""
    if not items:
        return {}

    system = _SYSTEM_PROMPT.format(
        doc_type=doc_type or "general",
        country=country or "generic",
    )
    payload = {"items": [{"id": _id, "text": text} for _id, text in items]}

    try:
        resp = _client.messages.create(
            model=MODEL,
            max_tokens=4096,
            system=system,
            messages=[{
                "role": "user",
                "content": (
                    "Translate the items below. Respond with ONLY the JSON object "
                    "matching the schema in the system prompt.\n\n"
                    + json.dumps(payload, ensure_ascii=False)
                ),
            }],
        )
    except Exception as e:
        log.error("API de traducción falló: %s", e)
        return {}

    raw = resp.content[0].text.strip() if resp.content else ""
    # Defensa contra Claude que a veces envuelve en ```json a pesar del prompt
    raw = re.sub(r'^```(?:json)?\s*', '', raw)
    raw = re.sub(r'\s*```$', '', raw)

    try:
        parsed = json.loads(raw)
    except json.JSONDecodeError as e:
        log.error("Respuesta del LLM no es JSON válido: %s; raw[:200]=%r", e, raw[:200])
        return {}

    out: dict[str, str] = {}
    for item in (parsed.get("items") or []):
        _id = item.get("id")
        text = item.get("text")
        if _id and isinstance(text, str):
            out[_id] = text
    return out


# ══════════════════════════════════════════════════════════════════════════════
# BANNER DE FALLO
# ══════════════════════════════════════════════════════════════════════════════
def _insert_failure_banner(doc):
    """Inserta un banner rojo al inicio del doc cuando la traducción falla."""
    try:
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        new = doc.add_paragraph()
        new.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = new.add_run(
            "⚠ ENGLISH TRANSLATION UNAVAILABLE — Document shown in original Spanish. "
            "Please try again later or contact support if the problem persists."
        )
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
        # Mover el párrafo al inicio del body manipulando el XML
        elem = new._element
        body = elem.getparent()
        body.remove(elem)
        body.insert(0, elem)
    except Exception as e:
        log.warning("No se pudo insertar el banner de fallo: %s", e)


# ══════════════════════════════════════════════════════════════════════════════
# API PÚBLICA
# ══════════════════════════════════════════════════════════════════════════════
def translate_doc_to_en(doc, doc_type: str = "general", cfg: dict | None = None):
    """Traduce el documento de español a inglés en sitio.

    Args:
        doc: instancia de docx.Document recién generada por un builder.
        doc_type: template_id (p. ej. "service_contract", "nda", "cv").
                  Se usa para dar contexto al LLM y elegir vocabulario legal.
        cfg: LEGAL_CONFIG del país (puede ser None). Se lee `cfg["pais"]`.

    Returns:
        El mismo objeto `doc`, traducido in place.

    Pipeline:
      1. Cargar caché desde disco (idempotente).
      2. Recolectar párrafos traducibles (body + tablas + headers + footers).
      3. Cada párrafo: hit en caché → no API, miss → cola para batch.
      4. Mandar lo no cacheado al LLM en batches de MAX_BATCH_SIZE.
      5. Sustituir texto in place, preservando formato del primer run.
      6. Persistir nuevas entradas al caché.

    Failure modes (todos non-fatal):
      - Sin API key → log warning, banner en el doc, doc devuelto en español.
      - API timeout/error → log error, banner, doc en español.
      - JSON malformado → log error, banner, doc en español.
    """
    _load_cache()
    cfg = cfg or {}
    country = cfg.get("pais", "generic")

    units = list(_iter_translatable_units(doc))
    if not units:
        return doc

    # Particionar en cache hits vs misses
    to_translate: list[tuple[str, str]] = []
    cached_hits: dict[str, str] = {}
    for loc, p in units:
        src = p.text
        hit = _cache.get(src)
        if hit is not None:
            cached_hits[loc] = hit
        else:
            to_translate.append((loc, src))

    log.info(
        "translate_doc_to_en: %d unidades, %d desde caché, %d a API",
        len(units), len(cached_hits), len(to_translate),
    )

    # Llamada(s) al LLM
    api_results: dict[str, str] = {}
    if to_translate:
        for batch in _chunks(to_translate, MAX_BATCH_SIZE):
            partial = _llm_translate(batch, doc_type, country)
            if not partial:
                # Un batch falló — degradación elegante
                _insert_failure_banner(doc)
                return doc
            api_results.update(partial)
            # Actualizar caché incrementalmente — si la app se mata a la mitad,
            # lo ya traducido queda persistido para la próxima.
            with _cache_lock:
                for loc, src in batch:
                    en = partial.get(loc)
                    if en:
                        _cache[src] = en
        _save_cache()

    # Aplicar traducciones in place
    for loc, p in units:
        en = cached_hits.get(loc) or api_results.get(loc)
        if en and en != p.text:
            _replace_paragraph_text(p, en)

    return doc


# ══════════════════════════════════════════════════════════════════════════════
# UTILIDADES OPCIONALES
# ══════════════════════════════════════════════════════════════════════════════
def cache_stats() -> dict:
    """Devuelve estadísticas del caché — útil para monitoreo / UI admin."""
    _load_cache()
    return {
        "entries": len(_cache),
        "path": CACHE_PATH,
        "size_bytes": os.path.getsize(CACHE_PATH) if os.path.exists(CACHE_PATH) else 0,
    }


def clear_cache():
    """Vacía el caché en memoria y borra el archivo. Útil para tests."""
    global _cache, _cache_loaded
    with _cache_lock:
        _cache = {}
        _cache_loaded = True
        try:
            if os.path.exists(CACHE_PATH):
                os.remove(CACHE_PATH)
        except Exception as e:
            log.warning("No se pudo borrar el caché: %s", e)
