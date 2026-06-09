"""
Tests parametrizados de smoke + regresión para los 38 builders de Word.

Estos tests NO requieren red, NO requieren API key real, NO requieren Mac.
Solo necesitan `python-docx`. Pensados para correr en CI (GitHub Actions).

Cubren tres clases de regresión:

  1. Smoke: cada template_id no crashea con un fixture mínimo.
  2. v15.4 regresión: secciones como list[dict] no crashea (atrapa el bug
     histórico de `sec.upper()` AttributeError en _build_executive_report).
  3. Calidad de output: el .docx generado existe, pesa razonable y contiene
     los datos críticos del usuario.

Ejecutar:
    pytest tests/test_word_builders.py -v
    pytest tests/test_word_builders.py -v -k nda      # solo el NDA
    pytest tests/test_word_builders.py --tb=short     # output compacto

v15.4 — nuevo en este release.
"""

import os
import sys
import pytest

# Permitir importar desde la raíz del proyecto sin instalar
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from office.word_controller import generate_word, WORD_TEMPLATES, _normalize_sections


# ── FIXTURES ──────────────────────────────────────────────────────────────────

@pytest.fixture
def fixture_datos():
    """Datos suficientes para que cualquier builder los use sin pegarse."""
    return {
        "pais": "MX",
        "empresa": "Empresa Test SA",
        # Partes (NDA, compraventa, etc.)
        "parte_a_nombre":  "Empresa A SA de CV",
        "parte_a_id":      "RFC AAA010101AAA",
        "parte_a_domicilio": "Av. Reforma 100, CDMX",
        "parte_b_nombre":  "Empresa B SA de CV",
        "parte_b_id":      "RFC BBB020202BBB",
        "parte_b_domicilio": "Av. Insurgentes 200, CDMX",
        "objeto":          "Negocio confidencial de test",
        # Laboral
        "patron_nombre":    "Empleadora SA de CV",
        "patron_id":        "RFC EMP030303CCC",
        "trabajador_nombre":"Juan Pérez García",
        "trabajador_id":    "CURP PEGA800101HDFRRN09",
        "puesto":           "Gerente de Operaciones",
        "salario":          "35000",
        "fecha_inicio":     "2026-06-01",
        "tipo_contrato":    "indefinido",
        # Comercial
        "monto":     "150000",
        "duracion":  "12 meses",
        "forma_pago":"Transferencia mensual",
        "cliente_nombre": "Cliente XYZ SA",
        "cliente_id":     "RFC CLI040404DDD",
        "emisor_nombre":  "Emisor SA de CV",
        "emisor_id":      "RFC EMI050505EEE",
        "items": [
            {"descripcion": "Consultoría estratégica", "cantidad": 10, "precio": 5000},
            {"descripcion": "Implementación piloto",   "cantidad": 1,  "precio": 50000},
        ],
        # Genérico
        "secciones": ["Sección uno", "Sección dos", "Sección tres"],
    }


# ── 1. SMOKE TEST — cada template_id no crashea ───────────────────────────────

@pytest.mark.parametrize("template_id", sorted(WORD_TEMPLATES.keys()))
def test_builder_smoke(template_id, fixture_datos, tmp_path, monkeypatch):
    """Cada builder debe generar un .docx sin crashear, con datos mínimos."""
    import office.word_controller as wc
    monkeypatch.setattr(wc, "OUTPUT_DIR", str(tmp_path))

    gen_data = {
        "template_id": template_id,
        "titulo": f"Smoke_{template_id}",
        "instrucciones": f"Documento de prueba para {template_id}",
        "datos": fixture_datos.copy(),
    }
    path = generate_word(gen_data)
    assert os.path.exists(path), f"Archivo no generado para {template_id}"
    assert os.path.getsize(path) > 1000, f"Archivo sospechosamente pequeño: {os.path.getsize(path)} bytes"
    assert path.endswith(".docx")


# ── 2. REGRESIÓN v15.4 — secciones como list[dict] no debe crashear ───────────

@pytest.mark.parametrize("template_id", sorted(WORD_TEMPLATES.keys()))
def test_builder_accepts_dict_sections(template_id, fixture_datos, tmp_path, monkeypatch):
    """v15.4 fix: secciones puede llegar como list[dict] desde el LLM.
    Atrapa el bug histórico de _build_executive_report (sec.upper() AttributeError).
    """
    import office.word_controller as wc
    monkeypatch.setattr(wc, "OUTPUT_DIR", str(tmp_path))

    datos = fixture_datos.copy()
    datos["secciones"] = [
        {"titulo": "Introducción", "contenido": "Texto de intro"},
        {"titulo": "Análisis de Riesgos", "contenido": "Texto de análisis"},
        {"titulo": "Recomendaciones", "contenido": "Texto de recos"},
    ]
    gen_data = {
        "template_id": template_id,
        "titulo": f"DictSec_{template_id}",
        "instrucciones": f"Test dict sections para {template_id}",
        "datos": datos,
    }
    path = generate_word(gen_data)
    assert os.path.exists(path)
    assert os.path.getsize(path) > 1000


# ── 3. _normalize_sections — tests unitarios del helper ───────────────────────

class TestNormalizeSections:
    """Tests directos del helper de normalización."""

    def test_list_of_strings(self):
        assert _normalize_sections(["A", "B", "C"]) == ["A", "B", "C"]

    def test_list_of_dicts_titulo(self):
        assert _normalize_sections(
            [{"titulo": "A"}, {"titulo": "B"}]
        ) == ["A", "B"]

    def test_list_of_dicts_title_en(self):
        assert _normalize_sections(
            [{"title": "A"}, {"title": "B"}]
        ) == ["A", "B"]

    def test_list_of_dicts_name(self):
        assert _normalize_sections([{"name": "A"}]) == ["A"]

    def test_mixed_str_and_dict(self):
        assert _normalize_sections(
            ["A", {"titulo": "B"}, {"name": "C"}]
        ) == ["A", "B", "C"]

    def test_empty_list(self):
        assert _normalize_sections([]) == []

    def test_none(self):
        assert _normalize_sections(None) == []

    def test_drops_empty_strings(self):
        assert _normalize_sections(["A", "", "  ", "B"]) == ["A", "B"]

    def test_drops_empty_dicts(self):
        assert _normalize_sections(
            [{}, {"titulo": ""}, {"titulo": "A"}, {"titulo": "   "}]
        ) == ["A"]

    def test_drops_none_items(self):
        assert _normalize_sections([None, "A", None, "B", None]) == ["A", "B"]

    def test_handles_unexpected_types(self):
        # Números o booleanos no deberían crashear
        result = _normalize_sections([42, True, "real"])
        assert "real" in result
        assert "42" in result

    def test_dict_priority_titulo_over_title(self):
        # Si tiene ambos, gana el ES (es la convención del proyecto)
        assert _normalize_sections(
            [{"titulo": "ES", "title": "EN"}]
        ) == ["ES"]


# ── 4. PRESERVACIÓN DE DATOS DEL USUARIO ──────────────────────────────────────

def test_user_data_preservation_nda(fixture_datos, tmp_path, monkeypatch):
    """Los datos críticos del usuario deben aparecer en el documento generado."""
    from docx import Document
    import office.word_controller as wc
    monkeypatch.setattr(wc, "OUTPUT_DIR", str(tmp_path))

    gen_data = {
        "template_id": "nda",
        "titulo": "NDA_Test",
        "datos": fixture_datos.copy(),
    }
    path = generate_word(gen_data)
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                text += "\n" + cell.text

    # Datos críticos que NO se pueden perder
    for critical in [fixture_datos["parte_a_nombre"], fixture_datos["parte_b_nombre"],
                     fixture_datos["parte_a_id"], fixture_datos["parte_b_id"]]:
        assert critical in text, f"Dato crítico '{critical}' no aparece en el documento"


# ── 5. COBERTURA DE PAÍSES ────────────────────────────────────────────────────

@pytest.mark.parametrize("pais", ["MX", "CO", "PE", "AR", "CL", "ES", "US",
                                    "BO", "EC", "UY", "PY", "VE", "PA", "GT",
                                    "HN", "NI", "SV", "CR", "CU", "DO"])
def test_employment_contract_all_countries(pais, fixture_datos, tmp_path, monkeypatch):
    """El contrato laboral debe generarse limpiamente en los 20 países hispanohablantes."""
    import office.word_controller as wc
    monkeypatch.setattr(wc, "OUTPUT_DIR", str(tmp_path))

    datos = fixture_datos.copy()
    datos["pais"] = pais
    gen_data = {
        "template_id": "employment_contract",
        "titulo": f"Contrato_{pais}",
        "datos": datos,
    }
    path = generate_word(gen_data)
    assert os.path.exists(path)

    # Verificar que no hay placeholders feos del fallback GENERIC
    from docx import Document
    text = "\n".join(p.text for p in Document(path).paragraphs)
    for t in Document(path).tables:
        for row in t.rows:
            for cell in row.cells:
                text += "\n" + cell.text

    # Estos placeholders solo deberían aparecer si el país no tiene config
    bad_placeholders = ["[País]", "[Moneda]", "[Ley Laboral aplicable del país]",
                        "[ID Fiscal Empresa]", "[Autoridad Fiscal]"]
    found = [p for p in bad_placeholders if p in text]
    assert not found, f"País {pais} muestra placeholders del fallback GENERIC: {found}"


# ── 6. CONFIG CENTRAL — modelo vigente ────────────────────────────────────────

def test_config_no_deprecated_model():
    """Ningún modelo default debe apuntar al modelo retirado claude-sonnet-4-20250514."""
    import config
    DEPRECATED = "claude-sonnet-4-20250514"
    assert config.DEFAULT_CHAT_MODEL       != DEPRECATED
    assert config.DEFAULT_VISION_MODEL     != DEPRECATED
    assert config.DEFAULT_TRANSLATOR_MODEL != DEPRECATED
    assert config.DEFAULT_LEGAL_MODEL      != DEPRECATED


def test_brain_uses_central_config():
    """brain.py debe consumir la config central, no hardcodear."""
    from agent import brain
    import config
    assert brain.MODEL_CHAT   == config.DEFAULT_CHAT_MODEL
    assert brain.MODEL_VISION == config.DEFAULT_VISION_MODEL
    assert brain.MODEL_TITLE  == config.DEFAULT_TITLE_MODEL
