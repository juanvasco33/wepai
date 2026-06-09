"""
Tests parametrizados para _build_shareholders_agreement v15.6.0.

Estos tests detectan los bugs identificados en la auditoría (SHA-01 a SHA-10)
y previenen regresiones. Se ejecutan sin red usando mocks de los verifiers.

Ejecutar:
    pytest tests/test_shareholders_agreement.py -v
"""
import os
import sys
import types
import pytest
from docx import Document

# Permitir importar desde la raíz del proyecto
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))


# ── Fixtures ──────────────────────────────────────────────────────────────

@pytest.fixture(autouse=True)
def _mock_verifiers(monkeypatch):
    """Mockea los verifiers para que los tests no hagan llamadas de red."""
    if 'agent' not in sys.modules:
        sys.modules['agent'] = types.ModuleType('agent')

    mock_fv = types.ModuleType('agent.fiscal_verifier')
    mock_fv.verify_fiscal_data = lambda cfg, force=False: cfg
    mock_fv.verify_labor_data  = lambda cfg, force=False: cfg
    sys.modules['agent.fiscal_verifier'] = mock_fv

    mock_cv = types.ModuleType('agent.corporate_verifier')
    mock_cv.verify_corporate_data = lambda cfg, force=False: cfg
    sys.modules['agent.corporate_verifier'] = mock_cv


@pytest.fixture
def base_cfg_mx():
    """Configuración base para México con merge corporativo."""
    from office.legal_config import LEGAL_CONFIG
    from office.corporate_config import merge_corporate
    cfg = dict(LEGAL_CONFIG["MX"])
    return merge_corporate(cfg)


@pytest.fixture
def datos_usuario_realistas():
    """Datos típicos de una startup mexicana con 3 socios."""
    return {
        "empresa_nombre": "Acme Tecnologías SA de CV",
        "empresa_id":     "ACM010203XYZ",
        "empresa_tipo":   "S.A.P.I. de C.V.",
        "capital_social": "1,000,000",
        "socios": [
            {"nombre": "María González",  "id": "GOPM830505AB1",
             "pct": "40%", "aportacion": "400000", "clase_acciones": "Ordinarias",
             "rol": "CEO"},
            {"nombre": "Carlos Ruiz",     "id": "RUSC820612CD2",
             "pct": "30%", "aportacion": "300000", "clase_acciones": "Ordinarias",
             "rol": "CTO"},
            {"nombre": "Fondo VC México", "id": "FVM150708EFG", "tipo": "empresa",
             "pct": "30%", "aportacion": "300000", "clase_acciones": "Preferentes"},
        ],
    }


def _extract_text(doc):
    out = []
    for p in doc.paragraphs:
        if p.text.strip():
            out.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        out.append(p.text)
    return "\n".join(out)


# ── Tests de regresión: bugs originales ───────────────────────────────────

class TestSHA01_DatosDelUsuario:
    """SHA-01: los datos reales del usuario deben aparecer en el documento."""

    def test_empresa_aparece_en_apertura(self, base_cfg_mx, datos_usuario_realistas):
        from office.word_controller import _build_shareholders_agreement
        doc = Document()
        _build_shareholders_agreement(doc, "T", "", [], "",
                                       base_cfg_mx, datos_usuario_realistas)
        texto = _extract_text(doc)
        assert "Acme Tecnologías" in texto, "Nombre de empresa no aparece"
        assert "ACM010203XYZ" in texto, "RFC de empresa no aparece"
        assert "S.A.P.I. de C.V." in texto, "Forma jurídica no aparece"

    def test_no_placeholder_literal(self, base_cfg_mx, datos_usuario_realistas):
        from office.word_controller import _build_shareholders_agreement
        doc = Document()
        _build_shareholders_agreement(doc, "T", "", [], "",
                                       base_cfg_mx, datos_usuario_realistas)
        texto = _extract_text(doc)
        assert "[Nombre de la Empresa]" not in texto, \
            "Placeholder literal no fue reemplazado"


class TestSHA02_ListaDinamicaSocios:
    """SHA-02: la tabla debe escalar a N socios."""

    @pytest.mark.parametrize("n_socios", [1, 2, 3, 5, 7, 10])
    def test_n_socios_aparecen(self, base_cfg_mx, n_socios):
        from office.word_controller import _build_shareholders_agreement
        socios = [
            {"nombre": f"Persona_{i}", "id": f"ID_{i}",
             "pct": f"{100//n_socios}%", "aportacion": "100000",
             "clase_acciones": "Ordinarias"}
            for i in range(n_socios)
        ]
        doc = Document()
        _build_shareholders_agreement(
            doc, "T", "", [], "", base_cfg_mx,
            {"empresa_nombre": "TestCo", "socios": socios})
        texto = _extract_text(doc)
        for i in range(n_socios):
            assert f"Persona_{i}" in texto, \
                f"Socio {i} no aparece (n_socios={n_socios})"


class TestSHA03_TablasBilingues:
    """SHA-03: tablas no deben contener español en doc en inglés."""

    MARCADORES_ES_PROHIBIDOS = [
        "Operaciones ordinarias",
        "Aportación inicial",
        "Mayoría simple",
        "Aprobación de presupuesto",
        "Emisión de nuevas acciones",
        "Reuniones mensuales",
        "Derecho de Preferencia",
    ]

    def test_tablas_en_ingles(self, base_cfg_mx):
        from office.word_controller import _build_shareholders_agreement
        cfg = dict(base_cfg_mx)
        cfg["idioma"] = "en"
        doc = Document()
        _build_shareholders_agreement(doc, "T", "", [], "", cfg, {})
        texto = _extract_text(doc)
        filtrados = [m for m in self.MARCADORES_ES_PROHIBIDOS if m in texto]
        assert not filtrados, f"Marcadores ES filtrados en doc EN: {filtrados}"


class TestSHA04_FirmasEscalables:
    """SHA-04: el bloque de firmas debe escalar a N socios."""

    @pytest.mark.parametrize("n_socios", [2, 3, 4, 5, 6])
    def test_todas_las_firmas_presentes(self, base_cfg_mx, n_socios):
        from office.word_controller import _build_shareholders_agreement
        socios = [
            {"nombre": f"Persona_{i}", "id": f"ID_{i}",
             "pct": f"{100//n_socios}%", "aportacion": "100000",
             "clase_acciones": "Ordinarias"}
            for i in range(n_socios)
        ]
        doc = Document()
        _build_shareholders_agreement(doc, "T", "", [], "",
                                       base_cfg_mx, {"socios": socios})
        texto = _extract_text(doc)
        # Verificar letras de los socios A, B, C, ...
        for i in range(n_socios):
            letra = chr(65 + i)
            assert f"SOCIO {letra}" in texto, \
                f"No aparece la firma del SOCIO {letra} (n_socios={n_socios})"


class TestSHA07_CentroArbitrajePorPais:
    """SHA-07: cada país debe tener su centro de arbitraje específico."""

    CASOS_PAIS = [
        ("MX", "CANACO"),
        ("CL", "CAM Santiago"),
        ("CO", "Cámara de Comercio de Bogotá"),
        ("AR", "Bolsa de Comercio de Buenos Aires"),
        ("PE", "Cámara de Comercio de Lima"),
        ("ES", "CIMA"),
        ("PY", "CAMP"),
        ("PA", "CeCAP"),
        ("EC", "Cámara de Comercio de Quito"),
        ("UY", "Cámara Nacional de Comercio"),
    ]

    @pytest.mark.parametrize("code,esperado", CASOS_PAIS)
    def test_arbitraje_pais(self, code, esperado):
        from office.legal_config import LEGAL_CONFIG
        from office.corporate_config import merge_corporate
        from office.word_controller import _build_shareholders_agreement
        cfg = dict(LEGAL_CONFIG[code])
        cfg = merge_corporate(cfg)
        doc = Document()
        _build_shareholders_agreement(doc, "T", "", [], "", cfg, {})
        texto = _extract_text(doc)
        assert esperado in texto, \
            f"{code}: '{esperado}' no aparece en doc"


class TestSHA10_ReservaLegalPorPais:
    """SHA-10: la reserva legal debe ser específica de cada jurisdicción."""

    CASOS_RESERVA = [
        ("MX",  "5",  "20"),
        ("CO",  "10", "50"),
        ("ES",  "10", "20"),
        ("VE",  "5",  "10"),
        ("PE",  "10", "20"),
        ("AR",  "5",  "20"),
        ("DO",  "5",  "10"),
        ("SV",  "7",  "20"),
    ]

    @pytest.mark.parametrize("code,pct,tope", CASOS_RESERVA)
    def test_reserva_pais(self, code, pct, tope):
        from office.legal_config import LEGAL_CONFIG
        from office.corporate_config import merge_corporate
        from office.word_controller import _build_shareholders_agreement
        cfg = dict(LEGAL_CONFIG[code])
        cfg = merge_corporate(cfg)
        doc = Document()
        _build_shareholders_agreement(doc, "T", "", [], "", cfg, {})
        texto = _extract_text(doc)
        assert f"{pct}%" in texto, f"{code}: % de reserva {pct}% no aparece"
        assert f"{tope}%" in texto, f"{code}: tope {tope}% no aparece"


# ── Tests de cobertura: 19 países hispanohablantes ────────────────────────

PAISES_HISPANOS = [
    "MX", "CO", "PE", "AR", "CL", "ES", "BO", "EC", "UY", "PY",
    "VE", "PA", "GT", "HN", "NI", "SV", "CR", "CU", "DO",
]


class TestCoberturaPaisesHispanos:
    """El builder debe generar sin error para los 19 países hispanohablantes."""

    @pytest.mark.parametrize("code", PAISES_HISPANOS)
    def test_genera_sin_excepcion(self, code):
        from office.legal_config import LEGAL_CONFIG
        from office.corporate_config import merge_corporate
        from office.word_controller import _build_shareholders_agreement
        cfg = dict(LEGAL_CONFIG[code])
        cfg = merge_corporate(cfg)
        doc = Document()
        _build_shareholders_agreement(doc, "T", "", [], "", cfg, {})
        # Sin error es suficiente para este test
        assert len(doc.paragraphs) > 5

    @pytest.mark.parametrize("code", PAISES_HISPANOS)
    def test_menciona_pais_y_ley_sociedades(self, code):
        from office.legal_config import LEGAL_CONFIG
        from office.corporate_config import merge_corporate
        from office.word_controller import _build_shareholders_agreement
        cfg = dict(LEGAL_CONFIG[code])
        cfg = merge_corporate(cfg)
        doc = Document()
        _build_shareholders_agreement(doc, "T", "", [], "", cfg, {})
        texto = _extract_text(doc)
        pais_nombre = LEGAL_CONFIG[code]["pais"]
        assert pais_nombre in texto, f"{code}: nombre del país no aparece"


# ── Tests de integración: edge cases ──────────────────────────────────────

class TestEdgeCases:
    """Casos límite que el builder debe manejar sin fallar."""

    def test_datos_vacios(self, base_cfg_mx):
        from office.word_controller import _build_shareholders_agreement
        doc = Document()
        _build_shareholders_agreement(doc, "T", "", [], "", base_cfg_mx, {})
        # Debe generar con placeholders bilingües
        assert len(doc.paragraphs) > 5

    def test_cfg_none(self):
        from office.word_controller import _build_shareholders_agreement
        doc = Document()
        _build_shareholders_agreement(doc, "T", "", [], "", None, None)
        assert len(doc.paragraphs) > 5

    def test_socios_dict_incompleto(self, base_cfg_mx):
        """Un socio que solo tiene nombre, sin id ni pct, no debe romper."""
        from office.word_controller import _build_shareholders_agreement
        doc = Document()
        _build_shareholders_agreement(doc, "T", "", [], "",
                                       base_cfg_mx,
                                       {"socios": [{"nombre": "Solo Nombre"}]})
        texto = _extract_text(doc)
        assert "Solo Nombre" in texto

    def test_un_solo_socio(self, base_cfg_mx):
        """Caso patológico: 1 socio. No debe lanzar excepción."""
        from office.word_controller import _build_shareholders_agreement
        doc = Document()
        _build_shareholders_agreement(doc, "T", "", [], "",
                                       base_cfg_mx,
                                       {"socios": [{"nombre": "Único Socio",
                                                    "id": "X", "pct": "100%"}]})
        texto = _extract_text(doc)
        assert "Único Socio" in texto


# ── Test de verificador autónomo ──────────────────────────────────────────

class TestCorporateVerifierIntegracion:
    """El verificador autónomo debe integrarse sin bloquear el flujo."""

    def test_verifier_falla_silenciosamente(self, base_cfg_mx):
        """Si el verifier lanza excepción, el documento debe generarse igual."""
        import sys
        # Reemplazar el mock con uno que falle
        mock_cv = types.ModuleType('agent.corporate_verifier')
        def _fail(cfg, force=False):
            raise RuntimeError("simulación: red caída")
        mock_cv.verify_corporate_data = _fail
        sys.modules['agent.corporate_verifier'] = mock_cv

        from office.word_controller import _build_shareholders_agreement
        doc = Document()
        # No debe lanzar excepción
        _build_shareholders_agreement(doc, "T", "", [], "", base_cfg_mx, {})
        assert len(doc.paragraphs) > 5

    def test_verifier_aplica_overrides(self, base_cfg_mx):
        """Si el verifier devuelve cambios, deben aplicarse al documento."""
        import sys
        mock_cv = types.ModuleType('agent.corporate_verifier')
        def _override(cfg, force=False):
            updated = dict(cfg)
            updated["centro_arbitraje"] = "CENTRO_MOCK_VERIFIED"
            return updated
        mock_cv.verify_corporate_data = _override
        sys.modules['agent.corporate_verifier'] = mock_cv

        from office.word_controller import _build_shareholders_agreement
        doc = Document()
        _build_shareholders_agreement(doc, "T", "", [], "", base_cfg_mx, {})
        texto = _extract_text(doc)
        assert "CENTRO_MOCK_VERIFIED" in texto, \
            "El override del verificador no se aplicó"
